import json
import os
import re
import sys
import io
import csv
import threading
import uuid
from datetime import datetime, timedelta
import subprocess
from pathlib import Path
from urllib.parse import urlparse
import logging
from logging.handlers import RotatingFileHandler

from functools import wraps

from flask import Flask, render_template, request, redirect, url_for, send_file, abort, flash
from flask_login import LoginManager, UserMixin, login_user, logout_user, current_user
from dotenv import load_dotenv
from sqlalchemy import text, bindparam
from docx import Document
from werkzeug.security import check_password_hash, generate_password_hash

from config.config import get_db_engine, require, init_context, get_settings
import config.config as _config_mod
from werkzeug.middleware.proxy_fix import ProxyFix

load_dotenv()

settings, paths = init_context()
logger = logging.getLogger("app")

def configure_logging(app, settings):
    log_path = settings.paths.log_dir / settings.log_file

    handler = RotatingFileHandler(
        log_path,
        maxBytes=settings.log_max_bytes,
        backupCount=settings.log_backup_count,
        encoding="utf-8",
    )

    level = getattr(logging, settings.log_level.upper(), logging.ERROR)
    handler.setLevel(level)

    handler.setFormatter(logging.Formatter(
        "%(asctime)s [%(levelname)s] %(name)s: %(message)s in %(pathname)s:%(lineno)d"
    ))

    # vyhneš sa duplikáciám handlerov (hlavne pri reloaderi)
    app.logger.handlers.clear()

    app.logger.setLevel(level)
    app.logger.addHandler(handler)

    # (voliteľné) zachytí aj root logy z knižníc
    root = logging.getLogger()
    root.setLevel(level)
    root.addHandler(handler)


app = Flask(__name__)
app.config["SECRET_KEY"] = require(settings.flask_secret_key, "FLASK_SECRET_KEY")
# za reverse proxy (Synology) – aby Flask vedel o pôvodnom HTTPS/hoste
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)
configure_logging(app, settings)


@app.template_filter("ts_human")
def ts_human(ts):
    """Unix timestamp → 'YYYY-MM-DD HH:MM'."""
    try:
        return datetime.fromtimestamp(int(ts)).strftime("%Y-%m-%d %H:%M")
    except Exception:
        return str(ts) if ts else "-"

# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------

_ROLE_RANK = {"user": 1, "power": 2, "admin": 3}


class User(UserMixin):
    def __init__(self, id: int, username: str, role: str):
        self.id = id
        self.username = username
        self.role = role

    def has_role(self, minimum: str) -> bool:
        return _ROLE_RANK.get(self.role, 0) >= _ROLE_RANK.get(minimum, 99)

    @property
    def is_admin(self) -> bool:
        return self.role == "admin"

    @property
    def is_power(self) -> bool:
        return self.has_role("power")


login_manager = LoginManager(app)
login_manager.login_view = "login"
login_manager.login_message = "Pre túto akciu sa musíš prihlásiť."
login_manager.login_message_category = "warn"


@login_manager.user_loader
def load_user(user_id: str):
    engine = get_db_engine()
    with engine.connect() as conn:
        row = conn.execute(
            text("SELECT id, username, role FROM users WHERE id = :id"),
            {"id": int(user_id)},
        ).fetchone()
    if row:
        return User(row[0], row[1], row[2])
    return None


def _role_required(minimum_role: str):
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if not current_user.is_authenticated:
                return login_manager.unauthorized()
            if not current_user.has_role(minimum_role):
                abort(403)
            return f(*args, **kwargs)
        return decorated
    return decorator


def user_required(f):
    return _role_required("user")(f)


def power_required(f):
    return _role_required("power")(f)


def admin_required(f):
    return _role_required("admin")(f)


# ---------------------------------------------------------------------------

SLOVAK_TERMS = settings.search_terms.get("slovakia", [])
TERM_RE = re.compile(r"(" + "|".join(re.escape(t) for t in SLOVAK_TERMS) + r")", re.IGNORECASE)


BUNDLE_GLOB = "news_bundle*.json"   # news_bundle.json (v run podadresároch)
PYTHON_BIN = sys.executable


def latest_bundle_path() -> str | None:
    p = max(paths.bundle_dir.rglob(BUNDLE_GLOB), default=None, key=lambda x: x.stat().st_mtime)
    return str(p) if p else None


def run_script(args: list[str], env: dict | None = None) -> tuple[int, str]:
    """
    Spustí skript a vráti (returncode, combined_output).
    """
    cp = subprocess.run(
        args,
        capture_output=True,
        text=True,
        env=env,
    )
    out = (cp.stdout or "") + ("\n" + cp.stderr if cp.stderr else "")
    return cp.returncode, out


def _job_path(job_id: str) -> Path:
    return paths.bundle_dir / "jobs" / f"{job_id}.json"


def _run_bg(job_id: str, args: list[str]) -> None:
    rc, out = run_script(args)
    result = {"status": "done" if rc == 0 else "error", "rc": rc, "out": out}
    p = _job_path(job_id)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps(result), encoding="utf-8")


def _job_status(job_id: str) -> dict:
    p = _job_path(job_id)
    if not p.exists():
        return {"status": "running", "rc": None, "out": ""}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {"status": "error", "rc": -1, "out": "Chyba čítania job súboru."}


def _parse_ingest_summary(out: str) -> dict | None:
    m = re.search(r'candidates=(\d+), deduped=(\d+), inserted=(\d+), updated=(\d+)', out)
    if m:
        return {"candidates": int(m.group(1)), "deduped": int(m.group(2)),
                "inserted": int(m.group(3)), "updated": int(m.group(4))}
    m = re.search(r'candidates=(\d+), deduped=(\d+)', out)
    if m:
        return {"candidates": int(m.group(1)), "deduped": int(m.group(2))}
    return None


def _parse_extract_summary(out: str) -> dict | None:
    m = re.search(r'ok=(\d+), soft_deleted=(\d+), commercial=(\d+), failed=(\d+)', out)
    if m:
        return {"ok": int(m.group(1)), "soft_deleted": int(m.group(2)),
                "commercial": int(m.group(3)), "failed": int(m.group(4))}
    return None


def split_sentences(text_: str):
    # Simple splitter good enough for FR/EN/AR news text
    return re.split(r'(?<=[.!?؟])\s+', text_)


def extract_context_sentences(text_: str, max_sentences: int = 3):
    if not text_:
        return []
    t = text_.replace("\n", " ")
    sentences = [s.strip() for s in split_sentences(t) if s.strip()]
    hits = []
    for s in sentences:
        if TERM_RE.search(s):
            hits.append(s)
        if len(hits) >= max_sentences:
            break
    return hits


def highlight_terms_html(s: str) -> str:
    # Very simple highlighting (safe enough if we escape later in template)
    return TERM_RE.sub(r"<mark>\1</mark>", s)


def build_filters(days: int, extraction: str, only_slovak: bool, relevance: str, include_deleted: bool, include_avoided: bool):
    """
    Zostaví WHERE klauzulu pre browse/export.

    extraction:      'ok' | 'unextracted' | 'all'
    relevance:       'all' | '1' | '0' | 'null' | '1_or_null' | '0_or_null'
    include_deleted: True  → zahrnie záznamy so soft-delete (deleted_at IS NOT NULL)
    include_avoided: True  → zahrnie záznamy zo zdrojov s is_avoided=1

    Tlačítká na browse.html:
      Reset        — days=0, extraction=all, rel=1_or_null, del=0, av=0  (default pri /browse)
      Všetky v DB  — days=0, extraction=all, rel=all,       del=1, av=1
      Skryté       — days=0, extraction=all, rel=0_or_null, del=1, av=1
    """
    where = []
    params = {}

    if days > 0:
        cutoff = datetime.now() - timedelta(days=days)
        where.append("a.last_seen_at >= :cutoff")
        params["cutoff"] = cutoff

    if not include_deleted:
        where.append("a.deleted_at IS NULL")

    if not include_avoided:
        where.append("s.is_avoided = 0")

    if extraction == "unextracted":
        where.append("a.content_text IS NULL")
    elif extraction == "ok":
        where.append("a.extraction_ok = 1")

    if relevance in ("1", "0"):
        where.append("a.relevance = :rel")
        params["rel"] = int(relevance)
    elif relevance == "null":
        where.append("a.relevance IS NULL")
    elif relevance == "1_or_null":
        where.append("(a.relevance = 1 OR a.relevance IS NULL)")
    elif relevance == "0_or_null":
        where.append("(a.relevance = 0 OR a.relevance IS NULL)")

    # Slovak context filter:
    # If extracted text exists -> search there; else fall back to title/snippet.
    if only_slovak:
        # Use LOWER + LIKE for portability
        like_terms = []
        for i, term in enumerate(SLOVAK_TERMS):
            key = f"t{i}"
            params[key] = f"%{term.lower()}%"
            like_terms.append(f"LOWER(COALESCE(a.content_text, a.title, a.snippet, '')) LIKE :{key}")
        where.append("(" + " OR ".join(like_terms) + ")")

    return (" AND ".join(where) if where else "1=1"), params


def browse_q_from_request():
    """
    Vytiahne len browse filtre (days/ok/sk/rel/del/av) z requestu.
    Funguje pre GET aj POST, lebo request.values = args + form.
    """
    allowed = {"days", "ok", "sk", "rel", "del", "av"}
    q = request.values.to_dict(flat=True)
    return {k: v for k, v in q.items() if k in allowed}


def get_selected_article_ids():
    """
    Robustne vytiahne vybrané article IDs z formu.
    Podporuje viacero názvov inputov: ids, article_id, article_ids, id, selected...
    """
    candidate_keys = ("ids", "article_id", "article_ids", "id", "selected", "selected_ids")

    raw = []
    for k in candidate_keys:
        raw.extend(request.form.getlist(k))

    # niekedy príde ako jeden CSV string: "1,2,3"
    if len(raw) == 1 and isinstance(raw[0], str) and "," in raw[0]:
        raw = [x.strip() for x in raw[0].split(",") if x.strip()]

    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            continue

    # deduplikácia, zachovať poradie
    seen = set()
    out = []
    for i in ids:
        if i not in seen:
            seen.add(i)
            out.append(i)
    return out


def normalize_domain(s: str) -> str:
    s = (s or "").strip().lower()
    if not s:
        return ""
    # ak niekto dá URL, zober netloc
    if "://" in s:
        try:
            s = urlparse(s).netloc.lower()
        except Exception:
            pass
    # odstráň path, query (ak zadal bez schémy)
    s = s.split("/")[0].split("?")[0].split("#")[0]
    if s.startswith("www."):
        s = s[4:]
    return s


@app.get("/")
def dashboard():
    engine = get_db_engine()
    with engine.begin() as conn:
        # posledná aktualizácia = posledný last_seen_at v tabuľke
        last_update = conn.execute(text("SELECT MAX(last_seen_at) FROM articles")).scalar()

        # zdroje (voliteľne – ak máš tabuľku sources)
        sources_count = conn.execute(text("SELECT COUNT(*) FROM sources")).scalar()

        # nájdené za 24h (podľa last_seen_at)
        found_24h = conn.execute(text("""
            SELECT COUNT(*) FROM articles
            WHERE last_seen_at >= (NOW() - INTERVAL 1 DAY)
              AND deleted_at IS NULL
        """)).scalar()

        # čaká na spracovanie (príklad: extraction_ok = 0 a nie je deleted a nie je z avoided source)
        pending = conn.execute(text("""
            SELECT COUNT(*)
            FROM articles a
            JOIN sources s ON s.id = a.source_id
            WHERE a.deleted_at IS NULL
              AND s.is_avoided = 0
              AND (a.extraction_ok = 0 OR a.extraction_ok IS NULL)
        """)).scalar()

        # chyby fetch/extract (len orientačne)
        errors = conn.execute(text("""
            SELECT COUNT(*)
            FROM articles
            WHERE deleted_at IS NULL
              AND fetch_error IS NOT NULL
              AND TRIM(fetch_error) <> ''
        """)).scalar()

    # formátovanie času na text (jednoducho, nech to v šablóne nie je None)
    last_update_str = ""
    if last_update:
        try:
            last_update_str = last_update.strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            last_update_str = str(last_update)

    return render_template(
        "index.html",              # NOVÁ dashboard šablóna
        last_update=last_update_str or "—",
        sources_count=sources_count or 0,
        found_24h=found_24h or 0,
        pending_count=pending or 0,
        errors_count=errors or 0,
        port=settings.flask_port,
    )


@app.get("/browse")
def browse():
    days = int(request.args.get("days", 0))
    extraction = request.args.get("extraction", "all")  # ok | unextracted | all
    only_slovak = request.args.get("sk", "0") == "1"
    relevance = request.args.get("rel", "1_or_null")  # all | 1 | 0 | null | 1_or_null
    include_deleted = request.args.get("del", "0") == "1"
    include_avoided = request.args.get("av", "0") == "1"

    where_sql, params = build_filters(days, extraction, only_slovak, relevance, include_deleted, include_avoided)

    limit = 300 if days > 0 or extraction == "all" else 100
    sql = f"""
        SELECT
            a.id,
            COALESCE(a.title_fr, a.title) AS title,
            s.domain,
            COALESCE(DATE_FORMAT(a.published_at_real, '%Y-%m-%d %H:%i:%s'), a.published_at_text) AS published,
            COALESCE(a.final_url, a.url) AS url,
            COALESCE(a.extraction_ok, 0) = 1 AS extraction_ok,
            a.fetch_error,
            LEFT(a.content_text, 4000) AS preview,
            LEFT(a.content_text_fr, 4000) AS preview_fr,
            a.relevance,
            a.relevance_note,
            a.deleted_at,
            s.id AS source_id,
            s.is_avoided,
            s.notes AS source_notes
        FROM articles a
        JOIN sources s ON s.id = a.source_id
        WHERE {where_sql}
        ORDER BY a.published_at_real IS NULL, a.published_at_real DESC
        LIMIT {limit}
    """

    engine = get_db_engine()
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()

    enriched = []
    for r in rows:
        preview = r.preview_fr or r.preview or ""
        ctx = extract_context_sentences(preview, max_sentences=3)
        ctx_hl = [highlight_terms_html(c) for c in ctx]
        enriched.append((r, ctx_hl))

    # Kompozícia výsledkov
    stats = {
        "extracted": sum(1 for r, _ in enriched if r.extraction_ok),
        "unextracted": sum(1 for r, _ in enriched if not r.extraction_ok),
        "rel_1": sum(1 for r, _ in enriched if r.relevance == 1),
        "rel_0": sum(1 for r, _ in enriched if r.relevance == 0),
        "rel_null": sum(1 for r, _ in enriched if r.relevance is None),
        "deleted": sum(1 for r, _ in enriched if r.deleted_at),
    }

    return render_template(
        "browse.html",
        rows=enriched,
        days=days,
        extraction=extraction,
        only_slovak=only_slovak,
        relevance=relevance,
        include_deleted=include_deleted,
        include_avoided=include_avoided,
        total=len(rows),
        limit=limit,
        stats=stats,
        port=settings.flask_port,
    )


@app.post("/bulk/delete")
@power_required
def bulk_delete():
    mode = request.form.get("mode", "soft")  # soft|hard
    ids_int = get_selected_article_ids()

    if not ids_int:
        logger.warning("bulk_delete called with no articles selected")
        return redirect(url_for("browse", **browse_q_from_request()))

    engine = get_db_engine()
    with engine.begin() as conn:
        if mode == "hard":
            conn.execute(text("""
                DELETE FROM run_articles
                WHERE article_id IN :ids
            """), {"ids": tuple(ids_int)})
            conn.execute(text("""
                DELETE FROM articles
                WHERE id IN :ids
            """), {"ids": tuple(ids_int)})
        else:
            conn.execute(text("""
                UPDATE articles
                SET deleted_at = NOW()
                WHERE id IN :ids
            """), {"ids": tuple(ids_int)})

    logger.info("bulk_delete (%s): %d articles %s", mode, len(ids_int), ids_int)
    return redirect(url_for("browse", **browse_q_from_request()))


@app.post("/undelete/<int:article_id>")
@power_required
def undelete(article_id: int):
    """Undo soft delete: vráti článok späť (deleted_at=NULL)."""
    engine = get_db_engine()
    with engine.begin() as conn:
        conn.execute(text("""
            UPDATE articles
            SET deleted_at = NULL
            WHERE id = :id
        """), {"id": article_id})

    logger.info("Article restored from soft delete: id=%s", article_id)
    return redirect(url_for("browse", **browse_q_from_request()))


@app.post("/bulk/exclude_domains")
@power_required
def bulk_exclude_domains():
    note = (request.form.get("note") or "").strip()[:255]
    ids_int = get_selected_article_ids()

    if not ids_int:
        return redirect(url_for("browse", **browse_q_from_request()))

    engine = get_db_engine()
    with engine.begin() as conn:
        # 1) zisti source_id pre vybrané články (EXPANDING!)
        q = text("""
            SELECT DISTINCT source_id
            FROM articles
            WHERE id IN :ids
        """).bindparams(bindparam("ids", expanding=True))

        source_rows = conn.execute(q, {"ids": ids_int}).fetchall()
        source_ids = [row[0] for row in source_rows if row[0] is not None]

        if not source_ids:
            logger.warning("bulk_exclude_domains: no source_ids found for articles %s", ids_int)
            return redirect(url_for("browse", **request.args))

        # 2) umlčať domény (EXPANDING!)
        u = text("""
            UPDATE sources
            SET is_avoided = 1
            WHERE id IN :sids
        """).bindparams(bindparam("sids", expanding=True))

        conn.execute(u, {"sids": source_ids})

        # 3) voliteľne uložiť dôvod do notes
        if note:
            # nastav iba tam, kde je notes prázdne; nech neprepisujeme ručne napísané dôvody
            u2 = text("""
                UPDATE sources
                SET notes = :note
                WHERE id IN :sids
                  AND (notes IS NULL OR TRIM(notes) = '')
            """).bindparams(bindparam("sids", expanding=True))

            conn.execute(u2, {"note": note, "sids": source_ids})

    logger.info("bulk_exclude_domains: %d sources muted from %d articles", len(source_ids), len(ids_int))
    return redirect(url_for("browse", **browse_q_from_request()))


@app.get("/stats")
def stats():
    """
    Štatistika umlčaných domén:
    - koľko článkov by sa defaultne skrylo (t.j. ušetrilo)
    - koľko z nich je soft-deleted
    - posledný výskyt
    + filter domain priamo v /stats
    + top 10 umlčaných podľa "saved"
    + (optional) články pre vybranú doménu
    """
    domain = (request.args.get("domain") or "").strip().lower()
    days = int(request.args.get("days", 0))
    cutoff = datetime(2000, 1, 1) if days == 0 else datetime.now() - timedelta(days=days)

    engine = get_db_engine()

    # 1) TOP 10 (vždy, bez domain filtra)
    with engine.begin() as conn:
        top10 = conn.execute(text("""
            SELECT
                s.id,
                s.domain,
                COALESCE(s.notes, '') AS notes,
                COUNT(a.id) AS total_articles,
                SUM(CASE WHEN a.deleted_at IS NULL THEN 1 ELSE 0 END) AS visible_if_unmuted,
                SUM(CASE WHEN a.deleted_at IS NOT NULL THEN 1 ELSE 0 END) AS soft_deleted,
                MAX(a.last_seen_at) AS last_seen
            FROM sources s
            JOIN articles a ON a.source_id = s.id
            WHERE s.is_avoided = 1
              AND a.last_seen_at >= :cutoff
            GROUP BY s.id, s.domain, s.notes
            ORDER BY visible_if_unmuted DESC, total_articles DESC
            LIMIT 10
        """), {"cutoff": cutoff}).mappings().all()

    # 1b) Global sum saved (všetky umlčané domény v okne)
    with engine.begin() as conn:
        total_saved_all = conn.execute(text("""
            SELECT
                COALESCE(SUM(CASE WHEN a.deleted_at IS NULL THEN 1 ELSE 0 END), 0) AS total_saved
            FROM sources s
            JOIN articles a ON a.source_id = s.id
            WHERE s.is_avoided = 1
              AND a.last_seen_at >= :cutoff
        """), {"cutoff": cutoff}).scalar_one()

    # 2) Tabuľka umlčaných domén (s voliteľným filtrom domain)
    where_extra = ""
    params = {"cutoff": cutoff}
    if domain:
        where_extra = " AND LOWER(s.domain) = :domain "
        params["domain"] = domain

    with engine.begin() as conn:
        rows = conn.execute(text(f"""
            SELECT
                s.id,
                s.domain,
                COALESCE(s.notes, '') AS notes,
                COUNT(a.id) AS total_articles,
                SUM(CASE WHEN a.deleted_at IS NULL THEN 1 ELSE 0 END) AS visible_if_unmuted,
                SUM(CASE WHEN a.deleted_at IS NOT NULL THEN 1 ELSE 0 END) AS soft_deleted,
                MAX(a.last_seen_at) AS last_seen
            FROM sources s
            JOIN articles a ON a.source_id = s.id
            WHERE s.is_avoided = 1
              AND a.last_seen_at >= :cutoff
              {where_extra}
            GROUP BY s.id, s.domain, s.notes
            ORDER BY visible_if_unmuted DESC, total_articles DESC
        """), params).mappings().all()

    # “ušetrené” = články z umlčaných domén, ktoré by sa inak zobrazili v UI defaultne
    enriched = []
    total_saved = 0
    for r in rows:
        saved = int(r["visible_if_unmuted"] or 0)
        total_saved += saved
        enriched.append({
            "id": r["id"],
            "domain": r["domain"],
            "notes": r["notes"],
            "total_articles": int(r["total_articles"] or 0),
            "saved": saved,
            "soft_deleted": int(r["soft_deleted"] or 0),
            "last_seen": r["last_seen"],
        })

    # 3) Optional: články pre vybranú doménu (aby klik “domain” malo okamžitý efekt priamo v /stats)
    articles = []
    if domain:
        with engine.begin() as conn:
            articles = conn.execute(text("""
                SELECT
                    a.id,
                    a.title,
                    COALESCE(DATE_FORMAT(a.published_at_real, '%%Y-%%m-%%d %%H:%%i'), a.published_at_text) AS published,
                    COALESCE(a.final_url, a.url) AS url,
                    a.extraction_ok,
                    a.fetch_error,
                    a.deleted_at,
                    a.last_seen_at
                FROM articles a
                JOIN sources s ON s.id = a.source_id
                WHERE LOWER(s.domain) = :domain
                  AND a.last_seen_at >= :cutoff
                ORDER BY a.last_seen_at DESC
                LIMIT 200
            """), {"domain": domain, "cutoff": cutoff}).mappings().all()

    return render_template(
        "stats.html",
        rows=enriched,
        top10=top10,
        articles=articles,
        days=days,
        domain=domain,
        total_saved=total_saved,
        total_saved_all=int(total_saved_all),
    )


@app.post("/label/<int:article_id>")
@power_required
def label(article_id: int):
    rel = request.form.get("relevance")  # "1" or "0" or "null"
    note = (request.form.get("note") or "").strip()[:255]

    if rel not in ("1", "0", "null"):
        abort(400)

    engine = get_db_engine()
    with engine.begin() as conn:
        if rel == "null":
            conn.execute(text("""
                UPDATE articles
                SET relevance=NULL, relevance_note=:note
                WHERE id=:id
            """), {"id": article_id, "note": note})
        else:
            conn.execute(text("""
                UPDATE articles
                SET relevance=:rel, relevance_note=:note
                WHERE id=:id
            """), {"id": article_id, "rel": int(rel), "note": note})

    logger.info("Article labeled: id=%s, relevance=%s, note=%r", article_id, rel, note)
    return redirect(url_for("browse", **browse_q_from_request()))


@app.post("/delete/<int:article_id>")
@power_required
def delete(article_id: int):
    mode = request.form.get("mode", "soft")  # soft | hard

    engine = get_db_engine()
    with engine.begin() as conn:
        if mode == "hard":
            # hard delete must remove links first
            conn.execute(text("DELETE FROM run_articles WHERE article_id=:id"), {"id": article_id})
            conn.execute(text("DELETE FROM articles WHERE id=:id"), {"id": article_id})
        else:
            conn.execute(text("UPDATE articles SET deleted_at=NOW() WHERE id=:id"), {"id": article_id})

    logger.info("Article deleted (%s): id=%s", mode, article_id)
    if mode == "soft":
        return redirect(url_for("article_detail", article_id=article_id))
    return redirect(url_for("browse", **browse_q_from_request()))


@app.get("/export/csv")
@power_required
def export_csv():
    days = int(request.args.get("days", 7))
    extraction = request.args.get("extraction", "ok")
    only_slovak = request.args.get("sk", "0") == "1"
    relevance = request.args.get("rel", "all")
    include_deleted = request.args.get("del", "0") == "1"
    include_avoided = request.args.get("av", "0") == "1"
    where_sql, params = build_filters(days, extraction, only_slovak, relevance, include_deleted, include_avoided)

    sql = f"""
        SELECT
            a.id,
            s.domain AS domain,
            COALESCE(DATE_FORMAT(a.published_at_real, '%Y-%m-%d %H:%i:%s'), a.published_at_text) AS published,
            COALESCE(a.final_url, a.url) AS url,
            a.title,
            a.snippet,
            a.relevance,
            a.relevance_note
        FROM articles a
        JOIN sources s ON s.id = a.source_id
        WHERE {where_sql}
        ORDER BY a.last_seen_at DESC
    """

    engine = get_db_engine()
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()

    out = io.StringIO()
    writer = csv.writer(out, delimiter=";")
    writer.writerow(["id", "domain", "published", "url", "title", "snippet", "relevance", "note"])

    if not rows:
        writer.writerow(["", "", "", "", "Žiadne nové správy", "", "", ""])
    else:
        for r in rows:
            writer.writerow([
                r.id, r.domain, r.published or "", r.url,
                r.title or "", r.snippet or "",
                "" if r.relevance is None else int(r.relevance),
                r.relevance_note or ""
            ])

    logger.info("CSV export: %d rows, %dd filter", len(rows), days)
    mem = io.BytesIO(out.getvalue().encode("utf-8-sig"))  # BOM helps Excel
    filename = f"dz_news_{days}d.csv"
    return send_file(mem, as_attachment=True, download_name=filename, mimetype="text/csv")


@app.get("/export/word")
@power_required
def export_word():
    days = int(request.args.get("days", 7))
    extraction = request.args.get("extraction", "ok")
    only_slovak = request.args.get("sk", "0") == "1"
    relevance = request.args.get("rel", "all")
    include_deleted = request.args.get("del", "0") == "1"
    include_avoided = request.args.get("av", "0") == "1"

    where_sql, params = build_filters(days, extraction, only_slovak, relevance, include_deleted, include_avoided)

    sql = f"""
        SELECT
            a.id,
            s.domain AS domain,
            COALESCE(DATE_FORMAT(a.published_at_real, '%Y-%m-%d %H:%i:%s'), a.published_at_text) AS published,
            COALESCE(a.final_url, a.url) AS url,
            a.title,
            a.snippet,
            a.content_text,
            a.relevance,
            a.relevance_note
        FROM articles a
        JOIN sources s ON s.id = a.source_id
        WHERE {where_sql}
        ORDER BY a.last_seen_at DESC
        LIMIT 200
    """

    engine = get_db_engine()
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()

    doc = Document()
    doc.add_heading(f"DZ News – monitoring (posledných {days} dní)", level=1)

    if not rows:
        doc.add_paragraph("Žiadne nové správy.")
    else:
        for r in rows:
            title = r.title or "— bez titulku —"
            doc.add_heading(title, level=2)
            meta = f"Zdroj: {r.domain} | Dátum: {r.published or ''}"
            if r.relevance is not None:
                meta += f" | Relevancia: {int(r.relevance)}"
            doc.add_paragraph(meta)
            doc.add_paragraph(r.url)

            if r.relevance_note:
                doc.add_paragraph(f"Poznámka: {r.relevance_note}")

            if r.snippet:
                doc.add_paragraph(r.snippet)

            # Slovak context sentences (strict, no inference)
            ctx = extract_context_sentences(r.content_text or "", max_sentences=3)
            if ctx:
                doc.add_paragraph("Výskyt SR v texte (explicitne):")
                for s in ctx:
                    doc.add_paragraph(f"- {s}")

    logger.info("Word export: %d rows, %dd filter", len(rows), days)
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)

    filename = f"dz_news_{days}d.docx"
    return send_file(mem, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/article/<int:article_id>/edit")
@admin_required
def article_edit(article_id: int):
    engine = get_db_engine()
    with engine.connect() as conn:
        r = conn.execute(text("""
            SELECT a.id, a.title, a.title_fr,
                   a.snippet, a.snippet_fr,
                   a.published_at_text, a.published_at_real, a.published_conf,
                   a.language, a.lang_detected,
                   a.relevance, a.relevance_note,
                   a.source_label,
                   a.content_text, a.content_text_fr,
                   a.extraction_ok,
                   s.domain
            FROM articles a
            JOIN sources s ON s.id = a.source_id
            WHERE a.id = :id
        """), {"id": article_id}).mappings().fetchone()
    if r is None:
        abort(404)
    return render_template("article_edit.html", r=r)


@app.post("/article/<int:article_id>/edit")
@admin_required
def article_edit_save(article_id: int):
    def _str(key):
        v = request.form.get(key, "").strip()
        return v if v else None

    def _dt(key):
        v = (request.form.get(key) or "").strip()
        if not v:
            return None
        try:
            return datetime.strptime(v, "%Y-%m-%d %H:%M:%S")
        except ValueError:
            try:
                return datetime.strptime(v, "%Y-%m-%d")
            except ValueError:
                return None

    published_conf = _str("published_conf")
    if published_conf not in (None, "search", "absolute"):
        published_conf = None

    relevance_raw = request.form.get("relevance", "null")
    relevance = None if relevance_raw == "null" else int(relevance_raw)

    extraction_ok = 1 if request.form.get("extraction_ok") == "1" else 0

    engine = get_db_engine()
    with engine.begin() as conn:
        conn.execute(text("""
            UPDATE articles SET
                title             = :title,
                title_fr          = :title_fr,
                snippet           = :snippet,
                snippet_fr        = :snippet_fr,
                published_at_text = :published_at_text,
                published_at_real = :published_at_real,
                published_conf    = :published_conf,
                language          = :language,
                lang_detected     = :lang_detected,
                relevance         = :relevance,
                relevance_note    = :relevance_note,
                source_label      = :source_label,
                content_text      = :content_text,
                content_text_fr   = :content_text_fr,
                extraction_ok     = :extraction_ok
            WHERE id = :id
        """), {
            "title":             _str("title"),
            "title_fr":          _str("title_fr"),
            "snippet":           _str("snippet"),
            "snippet_fr":        _str("snippet_fr"),
            "published_at_text": _str("published_at_text"),
            "published_at_real": _dt("published_at_real"),
            "published_conf":    published_conf,
            "language":          _str("language"),
            "lang_detected":     _str("lang_detected"),
            "relevance":         relevance,
            "relevance_note":    _str("relevance_note"),
            "source_label":      _str("source_label"),
            "content_text":      _str("content_text"),
            "content_text_fr":   _str("content_text_fr"),
            "extraction_ok":     extraction_ok,
            "id":                article_id,
        })
    logger.info("Article edited: id=%s by=%s", article_id, current_user.username)
    flash("Článok uložený.", "ok")
    return redirect(url_for("article_detail", article_id=article_id))


@app.get("/article/<int:article_id>")
def article_detail(article_id: int):
    engine = get_db_engine()
    with engine.begin() as conn:
        r = conn.execute(text("""
            SELECT
                a.id,
                a.source_id,
                s.domain,
                s.is_preferred,
                s.is_avoided,
                s.notes AS source_notes,

                a.url,
                a.final_url,
                a.final_url_canonical,
                a.final_url_hash,

                a.url_canonical,
                a.url_hash,

                a.title,
                a.published_at_text,
                a.published_at_real,
                a.published_conf,

                a.snippet,
                a.language,
                a.lang_detected,

                a.extraction_ok,
                a.source_label,

                a.first_seen_at,
                a.last_seen_at,
                a.fetched_at,
                a.http_status,
                a.fetch_error,

                a.content_text,
                a.content_text_fr,
                a.title_fr,
                a.snippet_fr,
                a.content_hash,

                a.ingestion_engine,
                a.ingestion_query_id,
                a.ingestion_rank,

                a.relevance,
                a.relevance_note,
                a.deleted_at
            FROM articles a
            JOIN sources s ON s.id = a.source_id
            WHERE a.id = :id
        """), {"id": article_id}).mappings().fetchone()

    if r is None:
        abort(404)

    # Preferuj final_url (ak existuje), inak url
    online_url = r["final_url"] or r["url"]

    published_display = (
        r["published_at_real"].strftime("%Y-%m-%d %H:%M:%S")
        if r["published_at_real"]
        else r["published_at_text"] or ""
    )

    ctx_plain_fr = []
    ctx_hl_fr = []

    try:
        ctx_plain = extract_context_sentences(
            r.get("content_text") or "",
            max_sentences=6
        )
        ctx_hl = [highlight_terms_html(s) for s in ctx_plain]
    except Exception:
        ctx_plain = []
        ctx_hl = []

    if r.get("content_text_fr"):
        try:
            ctx_plain_fr = extract_context_sentences(
                r.get("content_text_fr") or "",
                max_sentences=6
            )
            ctx_hl_fr = [highlight_terms_html(s) for s in ctx_plain_fr]
        except Exception:
            ctx_plain_fr = []
            ctx_hl_fr = []

    content_text_hl = highlight_terms_html(r.get("content_text") or "")
    snippet_hl = highlight_terms_html(r.get("snippet") or "")
    content_text_fr_hl = highlight_terms_html(r.get("content_text_fr") or "")
    snippet_fr_hl = highlight_terms_html(r.get("snippet_fr") or "")

    return render_template(
        "article.html",
        r=r,
        online_url=online_url,
        ctx_plain=ctx_plain,
        ctx_hl=ctx_hl,
        ctx_plain_fr=ctx_plain_fr,
        ctx_hl_fr=ctx_hl_fr,
        published_display=published_display,
        sk_context_found=len(ctx_hl) > 0,
        content_text_hl=content_text_hl,
        snippet_hl=snippet_hl,
        content_text_fr_hl=content_text_fr_hl,
        snippet_fr_hl=snippet_fr_hl,
    )


@app.post("/article/<int:article_id>/label")
@power_required
def article_label(article_id: int):
    relevance_raw = request.form.get("relevance", "null")
    note = (request.form.get("note") or "").strip()[:255]

    if relevance_raw == "null":
        relevance = None
    elif relevance_raw in ("0", "1"):
        relevance = int(relevance_raw)
    else:
        abort(400)

    engine = get_db_engine()
    with engine.begin() as conn:
        conn.execute(text("""
            UPDATE articles
            SET relevance = :relevance,
                relevance_note = :note
            WHERE id = :id
        """), {"relevance": relevance, "note": note if note else None, "id": article_id})

    logger.info("Article labeled: id=%s, relevance=%s, note=%r", article_id, relevance_raw, note)
    return redirect(url_for("article_detail", article_id=article_id))


@app.post("/article/<int:article_id>/exclude_domain")
@power_required
def article_exclude_domain(article_id: int):
    note = (request.form.get("note") or "").strip()[:255]  # voliteľné: dôvod

    engine = get_db_engine()
    with engine.begin() as conn:
        r = conn.execute(text("""
            SELECT source_id
            FROM articles
            WHERE id = :id
        """), {"id": article_id}).scalar()

        if not r:
            abort(404)

        conn.execute(text("""
            UPDATE sources
            SET is_avoided = 1
            WHERE id = :sid
        """), {"sid": r})

        if note:
            # nastav dôvod len ak prázdny
            conn.execute(text("""
                UPDATE sources
                SET notes = :note
                WHERE id = :sid AND (notes IS NULL OR TRIM(notes) = '')
            """), {"note": note, "sid": r})

    logger.info("Source muted from article: article_id=%s, source_id=%s", article_id, r)
    return redirect(url_for("article_detail", article_id=article_id))


@app.post("/article/<int:article_id>/unexclude_domain")
@power_required
def article_unexclude_domain(article_id: int):
    engine = get_db_engine()
    with engine.begin() as conn:
        sid = conn.execute(text("""
            SELECT source_id FROM articles WHERE id=:id
        """), {"id": article_id}).scalar()
        if not sid:
            abort(404)

        conn.execute(text("""
            UPDATE sources SET is_avoided = 0 WHERE id = :sid
        """), {"sid": sid})

    logger.info("Source unmuted from article: article_id=%s, source_id=%s", article_id, sid)
    return redirect(url_for("article_detail", article_id=article_id))


@app.post("/article/<int:article_id>/fetch")
@user_required
def fetch_extract_article(article_id: int):
    code, out = run_script([PYTHON_BIN, "refetch_article.py", "--article-id", str(article_id)])
    if code == 0:
        logger.info("Article fetch OK: id=%s", article_id)
        flash("Fetch & extract OK.", "success")
    else:
        logger.error("Article fetch failed: id=%s\n%s", article_id, out)
        flash("Fetch & extract FAIL — pozri log.", "error")

    return redirect(url_for("article_detail", article_id=article_id))


@app.post("/article/<int:article_id>/translate")
@user_required
def article_translate(article_id: int):
    engine = get_db_engine()
    with engine.begin() as conn:
        row = conn.execute(
            text("SELECT content_text, snippet, title, content_text_fr, snippet_fr, title_fr FROM articles WHERE id = :id"),
            {"id": article_id}
        ).fetchone()

    if not row or not row[0]:
        flash("Článok nemá extrahovaný text — preklad nie je možný.", "warn")
        return redirect(url_for("article_detail", article_id=article_id))

    api_key = settings.deepl_api_key
    if not api_key:
        flash("DEEPL_API_KEY nie je nastavený v .env.", "bad")
        return redirect(url_for("article_detail", article_id=article_id))

    texts, keys = [], []
    if not row[3]:  # content_text_fr
        texts.append(row[0]); keys.append("content_text_fr")
    if not row[4] and row[1]:  # snippet_fr
        texts.append(row[1]); keys.append("snippet_fr")
    if not row[5] and row[2]:  # title_fr
        texts.append(row[2]); keys.append("title_fr")

    if not texts:
        flash("Všetky polia sú už preložené.", "ok")
        return redirect(url_for("article_detail", article_id=article_id))

    try:
        from translate import translate_ar_fr
        translated = translate_ar_fr(api_key, texts)
        updates = dict(zip(keys, translated))
        updates["id"] = article_id
    except Exception as e:
        logger.error("DeepL translation failed: id=%s, error=%s", article_id, e)
        flash(f"Preklad zlyhal: {e}", "bad")
        return redirect(url_for("article_detail", article_id=article_id))

    set_clause = ", ".join(f"{k} = :{k}" for k in keys)
    with engine.begin() as conn:
        conn.execute(text(f"UPDATE articles SET {set_clause} WHERE id = :id"), updates)

    logger.info("Article translated AR→FR: id=%s", article_id)
    flash("Preklad AR→FR uložený.", "ok")
    return redirect(url_for("article_detail", article_id=article_id))


@app.post("/article/<int:article_id>/set_content")
@power_required
def article_set_content(article_id: int):
    content = (request.form.get("content_text") or "").strip()
    if not content:
        flash("Obsah je prázdny — nič neuložené.", "warn")
        return redirect(url_for("article_detail", article_id=article_id))

    engine = get_db_engine()
    with engine.begin() as conn:
        conn.execute(text("""
            UPDATE articles
            SET content_text   = :content,
                content_hash   = MD5(:content),
                extraction_ok  = 1,
                fetch_error    = NULL,
                fetched_at     = NOW()
            WHERE id = :id
        """), {"content": content, "id": article_id})

    logger.info("Article content set manually: id=%s, chars=%s", article_id, len(content))
    flash("Obsah uložený.", "ok")
    return redirect(url_for("article_detail", article_id=article_id))


@app.post("/bulk/restore")
@power_required
def bulk_restore():
    ids_int = get_selected_article_ids()
    if not ids_int:
        logger.warning("bulk_restore called with no articles selected")
        return redirect(url_for("browse", **browse_q_from_request()))

    engine = get_db_engine()
    with engine.begin() as conn:
        q = text("UPDATE articles SET deleted_at = NULL WHERE id IN :ids")
        conn.execute(q, {"ids": tuple(ids_int)})

    logger.info("bulk_restore: %d articles restored", len(ids_int))
    return redirect(url_for("browse", **browse_q_from_request()))


@app.get("/search")
@user_required
def search_page():
    lb = latest_bundle_path()
    return render_template("search.html", latest_bundle=lb)

def _bundle_articles(bundle_path: str | None) -> list[dict]:
    """Extrahuje všetky news_results z bundle JSON."""
    if not bundle_path:
        return []
    try:
        import json as _json
        data = _json.loads(Path(bundle_path).read_text(encoding="utf-8"))
        arts = []
        for resp in (data.get("responses_clean") or {}).values():
            if resp:
                arts.extend(resp.get("news_results") or [])
        return arts
    except Exception as e:
        logger.warning("_bundle_articles error: %s", e)
        return []


@app.post("/search/run")
@user_required
def search_run():
    cmd = [PYTHON_BIN, "search_flow_news.py"]
    hl = request.form.get("hl", "").strip()
    gl = request.form.get("gl", "").strip()
    num = request.form.get("num", "").strip()
    when = request.form.get("when", "").strip()
    if hl:
        cmd += ["--hl", hl]
    if gl:
        cmd += ["--gl", gl]
    if num:
        cmd += ["--num", num]
    if when:
        cmd += ["--when", when]
    code, out = run_script(cmd)
    lb = latest_bundle_path()
    if code == 0:
        logger.info("search_run OK: exit_code=%s, bundle=%s", code, lb)
        arts = _bundle_articles(lb)
        flash(f"OK: vyhľadávanie dokončené. Nájdených článkov: {len(arts)}", "ok")
    else:
        logger.error("search_run FAILED: exit_code=%s\n%s", code, out)
        arts = []
        flash("CHYBA: vyhľadávanie zlyhalo (pozri log).", "bad")
    return render_template("search.html", latest_bundle=lb, last_log=out, last_rc=code, bundle_articles=arts)


@app.post("/search/mfa-run")
@user_required
def search_mfa_run():
    when = request.form.get("when", "").strip() or "30d"
    code, out = run_script([PYTHON_BIN, "search_mfa_gov.py", "--when", when])
    lb = latest_bundle_path()
    mfa_articles = []
    if code == 0:
        mfa_articles = _bundle_articles(lb)
        flash(f"OK: MFA prehľadávanie dokončené. Nájdených: {len(mfa_articles)}", "ok")
    else:
        flash("CHYBA: MFA prehľadávanie zlyhalo (pozri log).", "bad")
    return render_template("search.html", latest_bundle=lb, last_log=out, last_rc=code,
                           mfa_articles=mfa_articles, mfa_when=when)


def _list_runs() -> list[dict]:
    """Vráti zoznam runov zoradených od najnovšieho."""
    import json as _json
    runs = []
    for run_dir in sorted(paths.runs_dir.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        bundle = run_dir / "news_bundle.json"
        if not bundle.exists():
            continue
        entry = {"run_id": run_dir.name, "bundle_path": str(bundle), "articles": 0, "timestamp": "", "hl": "", "gl": ""}
        try:
            data = _json.loads(bundle.read_text(encoding="utf-8"))
            rc = data.get("responses_clean") or {}
            entry["articles"] = sum(len((v or {}).get("news_results") or []) for v in rc.values())
            run_meta = data.get("run") or {}
            entry["timestamp"] = run_meta.get("timestamp", "")
            entry["hl"] = run_meta.get("hl", "")
            entry["gl"] = run_meta.get("gl", "")
        except Exception:
            pass
        runs.append(entry)
    return runs


@app.get("/process")
@user_required
def process_page():
    return render_template("process.html", runs=_list_runs(), latest_bundle=latest_bundle_path())


@app.post("/process/ingest_latest")
@user_required
def process_ingest_latest():
    lb = request.form.get("bundle_path") or latest_bundle_path()
    if not lb:
        logger.warning("process_ingest_latest: no bundle file found")
        flash("Nenašiel som žiadny bundle.", "bad")
        return redirect(url_for("process_page"))

    logger.info("process_ingest_latest started: bundle=%s", lb)
    code, out = run_script([PYTHON_BIN, "ingest_to_dz_news_reworked.py", lb])
    summary = _parse_ingest_summary(out)
    if code == 0:
        logger.info("process_ingest_latest OK: exit_code=%s, bundle=%s", code, lb)
        flash("OK: ingest hotový.", "ok")
    else:
        logger.error("process_ingest_latest FAILED: exit_code=%s\n%s", code, out)
        flash("CHYBA: ingest zlyhal (pozri log).", "bad")

    return render_template("process.html", runs=_list_runs(), latest_bundle=latest_bundle_path(),
                           last_log=out, last_rc=code, ingested_bundle=lb,
                           last_summary=summary, last_op="ingest")


@app.post("/process/extract_bulk")
@user_required
def process_extract_bulk():
    limit = request.form.get("limit", "50").strip() or "50"
    job_id = uuid.uuid4().hex[:12]
    logger.info("process_extract_bulk started: limit=%s, job=%s", limit, job_id)
    t = threading.Thread(
        target=_run_bg,
        args=(job_id, [PYTHON_BIN, "extract_bulk.py", "--limit", limit]),
        daemon=True,
    )
    t.start()
    return render_template("process.html", runs=_list_runs(), latest_bundle=latest_bundle_path(),
                           last_op="extract", extract_job_id=job_id)


@app.get("/process/job/<job_id>")
@user_required
def job_status(job_id: str):
    from flask import jsonify
    job = _job_status(job_id)
    if job["status"] == "done":
        job["summary"] = _parse_extract_summary(job.get("out", ""))
    return jsonify(job)


@app.get("/errors")
@power_required
def errors_page():
    days = int(request.args.get("days", 0))
    engine = get_db_engine()

    date_filter = "AND a.last_seen_at >= (NOW() - INTERVAL :days DAY)" if days > 0 else ""
    limit = 15 if days == 0 else 300
    params = {"days": days} if days > 0 else {}

    with engine.begin() as conn:
        rows = conn.execute(text(f"""
            SELECT a.id, a.title, s.domain,
                   COALESCE(DATE_FORMAT(a.published_at_real, '%%Y-%%m-%%d %%H:%%i'), a.published_at_text) AS published,
                   a.fetch_error, a.http_status, a.last_seen_at
            FROM articles a
            JOIN sources s ON s.id = a.source_id
            WHERE a.deleted_at IS NULL
              {date_filter}
              AND (a.fetch_error IS NOT NULL AND TRIM(a.fetch_error) <> '')
            ORDER BY a.last_seen_at DESC
            LIMIT {limit}
        """), params).mappings().all()

    return render_template("errors.html", rows=rows, days=days)



@app.get("/sources")
@admin_required
def sources_page():
    q = (request.args.get("q") or "").strip().lower()
    days = int(request.args.get("days", 30))
    only_avoided = (request.args.get("only_avoided", "0") == "1")
    cutoff = datetime.now() - timedelta(days=days)

    where = ["1=1"]
    params = {"cutoff": cutoff}

    if only_avoided:
        where.append("s.is_avoided = 1")

    if q:
        where.append("LOWER(s.domain) LIKE :q")
        params["q"] = f"%{q}%"

    sql = f"""
        SELECT
            s.id,
            s.domain,
            COALESCE(s.notes, '') AS notes,
            COALESCE(s.is_avoided, 0) AS is_avoided,

            COUNT(a.id) AS total_articles,
            SUM(CASE WHEN a.deleted_at IS NULL THEN 1 ELSE 0 END) AS visible_if_unmuted,
            SUM(CASE WHEN a.deleted_at IS NOT NULL THEN 1 ELSE 0 END) AS soft_deleted,
            MAX(a.last_seen_at) AS last_seen
        FROM sources s
        LEFT JOIN articles a
               ON a.source_id = s.id
              AND a.last_seen_at >= :cutoff
        WHERE {" AND ".join(where)}
        GROUP BY s.id, s.domain, s.notes, s.is_avoided
        ORDER BY s.is_avoided DESC, visible_if_unmuted DESC, total_articles DESC, s.domain ASC
    """

    engine = get_db_engine()
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).mappings().all()

    enriched = []
    for r in rows:
        saved = int(r["visible_if_unmuted"] or 0) if int(r["is_avoided"] or 0) == 1 else 0
        # Pozn.: "ušetrené" má význam najmä pri mute doménach; pri nemute dáme 0 (aby to nemiatlo).
        enriched.append({
            "id": r["id"],
            "domain": r["domain"],
            "notes": r["notes"],
            "is_avoided": int(r["is_avoided"] or 0),
            "total_articles": int(r["total_articles"] or 0),
            "soft_deleted": int(r["soft_deleted"] or 0),
            "saved": saved,
            "last_seen": r["last_seen"],
        })

    return render_template(
        "sources.html",
        rows=enriched,
        q=q,
        days=days,
        only_avoided=only_avoided,
    )


@app.post("/source/create")
@admin_required
def source_create():
    domain_in = request.form.get("domain", "")
    note = (request.form.get("note") or "").strip()

    domain = normalize_domain(domain_in)
    if not domain:
        flash("Chýba doména.", "warn")
        return redirect("/sources")

    engine = get_db_engine()
    with engine.begin() as conn:
        # upsert-like správanie: ak existuje, len update note (ak zadané)
        existing = conn.execute(
            text("SELECT id FROM sources WHERE LOWER(domain) = :d"),
            {"d": domain}
        ).scalar()

        if existing:
            if note:
                conn.execute(
                    text("UPDATE sources SET notes = :n WHERE id = :id"),
                    {"n": note, "id": existing}
                )
            logger.warning("source_create: domain already exists: %s", domain)
            flash(f"Doména už existuje: {domain}", "warn")
        else:
            conn.execute(
                text("""
                    INSERT INTO sources (domain, notes, is_avoided)
                    VALUES (:d, :n, 0)
                """),
                {"d": domain, "n": note}
            )
            logger.info("source_create: new source added: %s", domain)
            flash(f"Pridané: {domain}", "ok")

    return redirect("/sources")


@app.post("/source/note/<int:source_id>")
@admin_required
def source_note(source_id: int):
    note = (request.form.get("note") or "").strip()

    engine = get_db_engine()
    with engine.begin() as conn:
        conn.execute(
            text("UPDATE sources SET notes = :n WHERE id = :id"),
            {"n": note if note else None, "id": source_id}
        )

    logger.info("Source note updated: id=%s, note=%r", source_id, note)
    flash("Poznámka uložená.", "ok")
    # vráť sa tam, kde bol user (ak posielal zo /sources alebo /stats)
    return redirect(request.referrer or "/sources")


@app.post("/source/avoid/<int:source_id>")
@admin_required
def source_avoid(source_id: int):
    engine = get_db_engine()
    with engine.begin() as conn:
        conn.execute(
            text("UPDATE sources SET is_avoided = 1 WHERE id = :id"),
            {"id": source_id}
        )
    logger.info("Source muted: id=%s", source_id)
    flash("Doména umlčaná (mute).", "ok")
    return redirect(request.referrer or "/sources")


@app.post("/source/unavoid/<int:source_id>")
@admin_required
def source_unavoid(source_id: int):
    engine = get_db_engine()
    with engine.begin() as conn:
        conn.execute(
            text("UPDATE sources SET is_avoided = 0 WHERE id = :id"),
            {"id": source_id}
        )
    logger.info("Source unmuted: id=%s", source_id)
    flash("Doména od-umlčaná (unmute).", "ok")
    return redirect(request.referrer or "/sources")


_VALID_ROLES = ("user", "power", "admin")


@app.get("/users")
@admin_required
def users_page():
    engine = get_db_engine()
    with engine.connect() as conn:
        rows = conn.execute(text("""
            SELECT id, username, role, created_at
            FROM users
            ORDER BY role DESC, username
        """)).mappings().fetchall()
    return render_template("users.html", users=rows)


@app.post("/user/create")
@admin_required
def user_create():
    username = (request.form.get("username") or "").strip()
    password = request.form.get("password") or ""
    role = request.form.get("role", "user")

    if not username or not password:
        flash("Username aj heslo sú povinné.", "warn")
        return redirect(url_for("users_page"))
    if role not in _VALID_ROLES:
        flash("Neplatná rola.", "warn")
        return redirect(url_for("users_page"))

    engine = get_db_engine()
    with engine.begin() as conn:
        existing = conn.execute(
            text("SELECT id FROM users WHERE username = :u"), {"u": username}
        ).fetchone()
        if existing:
            flash(f"Používateľ '{username}' už existuje.", "warn")
            return redirect(url_for("users_page"))
        conn.execute(
            text("INSERT INTO users (username, password_hash, role) VALUES (:u, :h, :r)"),
            {"u": username, "h": generate_password_hash(password), "r": role},
        )
    logger.info("User created: username=%s role=%s by=%s", username, role, current_user.username)
    flash(f"Používateľ '{username}' ({role}) vytvorený.", "ok")
    return redirect(url_for("users_page"))


@app.post("/user/<int:user_id>/role")
@admin_required
def user_set_role(user_id: int):
    role = request.form.get("role", "")
    if role not in _VALID_ROLES:
        flash("Neplatná rola.", "warn")
        return redirect(url_for("users_page"))
    if user_id == current_user.id:
        flash("Nemôžeš zmeniť vlastnú rolu.", "warn")
        return redirect(url_for("users_page"))
    engine = get_db_engine()
    with engine.begin() as conn:
        conn.execute(
            text("UPDATE users SET role = :r WHERE id = :id"),
            {"r": role, "id": user_id},
        )
    logger.info("User role changed: id=%s role=%s by=%s", user_id, role, current_user.username)
    flash("Rola zmenená.", "ok")
    return redirect(url_for("users_page"))


@app.post("/user/<int:user_id>/delete")
@admin_required
def user_delete(user_id: int):
    if user_id == current_user.id:
        flash("Nemôžeš zmazať vlastný účet.", "warn")
        return redirect(url_for("users_page"))
    engine = get_db_engine()
    with engine.begin() as conn:
        row = conn.execute(
            text("SELECT username FROM users WHERE id = :id"), {"id": user_id}
        ).fetchone()
        if not row:
            flash("Používateľ neexistuje.", "warn")
            return redirect(url_for("users_page"))
        conn.execute(text("DELETE FROM users WHERE id = :id"), {"id": user_id})
    logger.info("User deleted: id=%s username=%s by=%s", user_id, row[0], current_user.username)
    flash(f"Používateľ '{row[0]}' zmazaný.", "ok")
    return redirect(url_for("users_page"))


@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        engine = get_db_engine()
        with engine.connect() as conn:
            row = conn.execute(
                text("SELECT id, username, password_hash, role FROM users WHERE username = :u"),
                {"u": username},
            ).fetchone()
        if row and check_password_hash(row[2], password):
            user = User(row[0], row[1], row[3])
            login_user(user)
            next_page = request.args.get("next")
            return redirect(next_page or url_for("dashboard"))
        flash("Nesprávne meno alebo heslo.", "error")
    return render_template("login.html")


@app.post("/logout")
def logout():
    logout_user()
    return redirect(url_for("dashboard"))


# ---------------------------------------------------------------------------
# Admin – Queries
# ---------------------------------------------------------------------------

_VALID_QTYPES = ("q1", "q2", "q3")


def _read_queries() -> dict:
    try:
        return json.loads(paths.queries_path.read_text(encoding="utf-8"))
    except Exception:
        return {"active": {"q1": "default", "q2": "default", "q3": "default"},
                "q1_presets": [], "q2_presets": [], "q3_presets": []}


def _write_queries(data: dict):
    paths.queries_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _compute_default_queries() -> dict:
    """Vráti dynamicky zostavené default query stringy (rovnaká logika ako search_flow_news)."""
    sk = settings.search_terms.get("slovakia", [])
    dz = settings.search_terms.get("algeria", [])
    preferred = list(settings.preferred_domains)
    sk_part = "(" + " OR ".join(sk) + ")" if sk else ""
    dz_part = "(" + " OR ".join(f'"{t}"' if " " in t else t for t in dz) + ")" if dz else ""
    url_f = "-inurl:recherche -inurl:search -inurl:tag -inurl:tags -inurl:page"
    sites_q1 = "(" + " OR ".join(f"site:{d}" for d in sorted(preferred)) + ")" if preferred else "site:.dz"
    return {
        "q1": f"{sites_q1} {sk_part} {url_f}".strip(),
        "q2": f"site:.dz {sk_part} {url_f}".strip(),
        "q3": f"{sk_part} {dz_part}".strip(),
    }


@app.get("/admin/queries")
@admin_required
def queries_page():
    data = _read_queries()
    defaults = _compute_default_queries()
    return render_template("admin_queries.html", data=data, defaults=defaults)


@app.post("/admin/queries/activate")
@admin_required
def queries_activate():
    qtype = request.form.get("qtype", "").strip()
    preset_id = request.form.get("preset_id", "").strip()
    if qtype not in _VALID_QTYPES or not preset_id:
        flash("Neplatný vstup.", "bad")
        return redirect(url_for("queries_page"))
    data = _read_queries()
    ids = [p["id"] for p in data.get(f"{qtype}_presets", [])]
    if preset_id not in ids:
        flash("Preset neexistuje.", "bad")
        return redirect(url_for("queries_page"))
    data["active"][qtype] = preset_id
    _write_queries(data)
    flash(f"{qtype.upper()} → aktivovaný preset: {preset_id}", "ok")
    return redirect(url_for("queries_page"))


@app.post("/admin/queries/save")
@admin_required
def queries_save():
    qtype = request.form.get("qtype", "").strip()
    label = request.form.get("label", "").strip()
    query = request.form.get("query", "").strip()
    preset_id = request.form.get("preset_id", "").strip()  # prázdne = nový
    if qtype not in _VALID_QTYPES or not label or not query:
        flash("Vyplň label aj query.", "bad")
        return redirect(url_for("queries_page"))
    data = _read_queries()
    presets = data.setdefault(f"{qtype}_presets", [])
    if preset_id and preset_id != "default":
        # update existujúceho
        for p in presets:
            if p["id"] == preset_id:
                p["label"] = label
                p["query"] = query
                flash(f"Preset '{label}' aktualizovaný.", "ok")
                break
        else:
            flash("Preset nenájdený.", "bad")
            return redirect(url_for("queries_page"))
    else:
        # nový preset
        new_id = re.sub(r"[^a-z0-9_]", "_", label.lower())[:40]
        # ak id koliduje, pridaj suffix
        existing_ids = {p["id"] for p in presets}
        base_id = new_id
        i = 2
        while new_id in existing_ids:
            new_id = f"{base_id}_{i}"
            i += 1
        presets.append({"id": new_id, "label": label, "query": query})
        flash(f"Nový preset '{label}' uložený.", "ok")
    _write_queries(data)
    return redirect(url_for("queries_page"))


@app.post("/admin/queries/delete")
@admin_required
def queries_delete():
    qtype = request.form.get("qtype", "").strip()
    preset_id = request.form.get("preset_id", "").strip()
    if qtype not in _VALID_QTYPES or not preset_id or preset_id == "default":
        flash("Default preset nie je možné odstrániť.", "warn")
        return redirect(url_for("queries_page"))
    data = _read_queries()
    presets = data.get(f"{qtype}_presets", [])
    data[f"{qtype}_presets"] = [p for p in presets if p["id"] != preset_id]
    if data["active"].get(qtype) == preset_id:
        data["active"][qtype] = "default"
    _write_queries(data)
    flash(f"Preset '{preset_id}' odstránený.", "ok")
    return redirect(url_for("queries_page"))


# ---------------------------------------------------------------------------
# Admin – Search terms
# ---------------------------------------------------------------------------

_VALID_GROUPS = ("slovakia", "algeria")


def _reload_search_terms():
    """Prepíše SLOVAK_TERMS a TERM_RE po zmene search_terms.json."""
    global SLOVAK_TERMS, TERM_RE
    _config_mod._CACHED = None
    new_settings = get_settings(force_reload=True)
    SLOVAK_TERMS = new_settings.search_terms.get("slovakia", [])
    if SLOVAK_TERMS:
        TERM_RE = re.compile(
            r"(" + "|".join(re.escape(t) for t in SLOVAK_TERMS) + r")",
            re.IGNORECASE,
        )
    else:
        TERM_RE = re.compile(r"(?!)")


def _read_search_terms() -> dict:
    try:
        return json.loads(paths.search_terms_path.read_text(encoding="utf-8"))
    except Exception:
        return {"slovakia": [], "algeria": []}


def _write_search_terms(data: dict):
    paths.search_terms_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


@app.get("/admin/search-terms")
@admin_required
def search_terms_page():
    data = _read_search_terms()
    sk_terms = data.get("slovakia", [])
    dz_terms = data.get("algeria", [])
    # Query preview
    sk_q = " OR ".join(sk_terms) if sk_terms else "—"
    dz_q = " OR ".join(dz_terms) if dz_terms else "—"
    return render_template(
        "admin_search_terms.html",
        sk_terms=sk_terms,
        dz_terms=dz_terms,
        sk_q=sk_q,
        dz_q=dz_q,
    )


@app.post("/admin/search-terms/add")
@admin_required
def search_terms_add():
    group = request.form.get("group", "").strip()
    term = request.form.get("term", "").strip()
    if group not in _VALID_GROUPS or not term:
        flash("Neplatný vstup.", "bad")
        return redirect(url_for("search_terms_page"))
    data = _read_search_terms()
    lst = data.setdefault(group, [])
    if term not in lst:
        lst.append(term)
        _write_search_terms(data)
        _reload_search_terms()
        flash(f"Pridaný termín: {term}", "ok")
    else:
        flash("Termín už existuje.", "warn")
    return redirect(url_for("search_terms_page"))


@app.post("/admin/search-terms/delete")
@admin_required
def search_terms_delete():
    group = request.form.get("group", "").strip()
    term = request.form.get("term", "").strip()
    if group not in _VALID_GROUPS or not term:
        flash("Neplatný vstup.", "bad")
        return redirect(url_for("search_terms_page"))
    data = _read_search_terms()
    lst = data.get(group, [])
    if term in lst:
        lst.remove(term)
        _write_search_terms(data)
        _reload_search_terms()
        flash(f"Odstránený termín: {term}", "ok")
    return redirect(url_for("search_terms_page"))


if __name__ == "__main__":
    port = settings.flask_port
    app.run(host="127.0.0.1", port=port, debug=True)
